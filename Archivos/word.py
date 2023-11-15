#pip install python-docx

from docx import Document

doc = Document()
doc.add_paragraph('Hola, este es un texto de ejemplo en un documento Word.')
doc.save('ejemplo.docx')

doc = Document('ejemplo.docx')
for para in doc.paragraphs:
    print(para.text)